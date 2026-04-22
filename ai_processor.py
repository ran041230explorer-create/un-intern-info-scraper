"""
使用 Google Gemini API（google-genai SDK）对联合国岗位原始抓取结果进行深度清洗与翻译，
采用「Flash 主处理 + Flash-Lite 补救」的阶梯式策略，输出符合人工整理习惯的中文结构化 JSON。
"""

from __future__ import annotations

import asyncio
import json
import os
import re
import sys
import time
import traceback
from datetime import datetime
from typing import Any

from dotenv import load_dotenv
from google import genai
from google.genai import errors as genai_errors
from google.genai import types

load_dotenv()

# ---------------------------------------------------------------------------
# 代理：GitHub Actions 等 CI 不配置；本地从 .env / 环境变量读取（勿硬编码 URL）
# ---------------------------------------------------------------------------
ACTIVE_HTTP_PROXY_URL: str | None = None


def is_github_actions_environment() -> bool:
    """当 ``GITHUB_ACTIONS`` 存在且非空时视为云端 CI，不注入代理。"""
    return bool(os.getenv("GITHUB_ACTIONS", "").strip())


def configure_runtime_networking() -> None:
    """
    若当前为 CI（GITHUB_ACTIONS），不设置任何代理。
    否则在 ``load_dotenv()`` 之后读取 ``UN_HTTP_PROXY`` / ``HTTP_PROXY`` / ``HTTPS_PROXY``，
    若存在则写入进程环境并供 google-genai 的 httpx 客户端使用。
    """
    global ACTIVE_HTTP_PROXY_URL
    load_dotenv()
    if is_github_actions_environment():
        ACTIVE_HTTP_PROXY_URL = None
        return
    url = (
        (os.getenv("UN_HTTP_PROXY") or "").strip()
        or (os.getenv("HTTP_PROXY") or "").strip()
        or (os.getenv("HTTPS_PROXY") or "").strip()
    )
    ACTIVE_HTTP_PROXY_URL = url or None
    if ACTIVE_HTTP_PROXY_URL:
        os.environ["HTTP_PROXY"] = ACTIVE_HTTP_PROXY_URL
        os.environ["HTTPS_PROXY"] = ACTIVE_HTTP_PROXY_URL
        os.environ["http_proxy"] = ACTIVE_HTTP_PROXY_URL
        os.environ["https_proxy"] = ACTIVE_HTTP_PROXY_URL


# ---------------------------------------------------------------------------
# API：密钥由 .env（推荐）或系统环境变量 GEMINI_API_KEY 提供（勿在代码中硬编码）
# ---------------------------------------------------------------------------
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

MODEL_FLASH = "gemini-2.5-flash"
MODEL_FLASH_LITE = "gemini-2.5-flash-lite"

INPUT_JSON_PATH = os.getenv("INPUT_JSON_PATH", "test_results.json")
OUTPUT_JSON_PATH = os.getenv("OUTPUT_JSON_PATH", "final_chinese_results.json")
OUTPUT_DOCX_PATH = os.getenv(
    "OUTPUT_DOCX_PATH", "联合国实习岗位半月汇总.docx"
)

# Word 正文展示顺序（不含 processed_by）
DOCX_FIELD_KEYS: list[str] = [
    "职位",
    "组织",
    "学历要求",
    "语言要求",
    "时长",
    "申请截止时间",
    "工作地点",
    "详情请见",
]

_FULLWIDTH_DIGITS = str.maketrans("０１２３４５６７８９", "0123456789")

# 每条岗位请求之间的节流（秒）
THROTTLE_SECONDS = int(os.getenv("GEMINI_THROTTLE_SECONDS", "15"))

# 第一阶段每批处理的原始条数（充分利用 Flash 配额；整文件按该步长切块）
PHASE1_CHUNK_SIZE = int(os.getenv("GEMINI_PHASE1_CHUNK_SIZE", "20"))

# 强制 JSON 输出的 8 个中文键（original_url → 「详情请见」；deadline → 「申请截止时间」）
JOB_OUTPUT_JSON_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "职位": {"type": "string"},
        "组织": {"type": "string"},
        "学历要求": {"anyOf": [{"type": "string"}, {"type": "null"}]},
        "语言要求": {"type": "string"},
        "时长": {"type": "string"},
        "申请截止时间": {"type": "string"},
        "工作地点": {"type": "string"},
        "详情请见": {"type": "string"},
    },
    "required": [
        "职位",
        "组织",
        "学历要求",
        "语言要求",
        "时长",
        "申请截止时间",
        "工作地点",
        "详情请见",
    ],
}

# =============================================================================
# AI 处理规范（核心指令 / Prompt 模板）
# =============================================================================
SYSTEM_INSTRUCTION = """你是联合国招聘信息的资深编辑，擅长将英文岗位说明整理为简洁、准确的中文摘要。
你必须严格依据用户提供的「单条岗位原始 JSON」作答，不得编造事实。
输出必须是单个 JSON 对象（不要 Markdown 代码块、不要前后解释文字）。

对键「学历要求」的字符串值，必须遵守用户模板中的分行与 (a)(b)(c) 条款格式；将复杂准入条件拆成多条独立行。
除末行外每行末尾须为「；或」；末行以句号收尾。字符串中必须包含真实的换行（在 JSON 里用 \\n 转义表示），
不得把多条要求压成单行或仅用空格拼接。"""

USER_PROMPT_TEMPLATE = """请根据以下「联合国岗位原始数据」（JSON），按规则生成**一个** JSON 对象。

## 原始数据（单条岗位）
```json
{raw_job_json}
```

## 字段生成规则（必须遵守）

1. **职位**（键名：`职位`）
   - 格式：`[英文职位名] [中文职位名]`，中间**一个空格**，**不要**括号。
   - 英文取自 `job_title`；中文为准确、简洁的译名。

2. **组织**（键名：`组织`）
   - 格式：`[英文全称] [官方中文全称]`，中间**一个空格**。
   - 英文部门/机构以 `gray_box_details.department_office` 为主，可结合 `full_description_block` 中出现的上级机构名称补全为「常用英文全称」。
   - 中文须使用联合国系统内**通行/官方**译法（例如：United Nations Environment Programme → 联合国环境规划署）。

3. **学历要求**（键名：`学历要求`）
   - **仅从** `full_description_block` 中提取与学历、在读状态、学位相关的**硬性要求**；按准入逻辑拆成多条，**不得**编造原文未给出的条件。
   - **格式（必须严格遵守）**
     - 引导词：整条字符串**必须以** `学历要求：` 开头（全角冒号）。
     - 条款拆分：各并列/选择条件依次用 **(a)**、**(b)**、**(c)** …（小写字母 + 半角括号）编号。
     - **强制换行**：每个条款**单独一行**（字符串内使用 JSON 的 `\\n` 表示换行，解析后即为真实换行符）。
     - **连接词**：除**最后一条**外，每一行末尾必须紧跟 **`；或`**（全角分号 + 汉字「或」）；**最后一条**行末用句号 **`。`** 收尾，**不得**再写「；或」。
   - **Few-shot（结构与换行示范，内容需按原文调整）**
```
学历要求：(a) 已注册研究生课程（第二大学学位或同等学历，或更高）；或
(b) 已注册第一个大学学位课程（至少学士学位或同等学历）的最后一年；或
(c) 已获得大学学位（学士、硕士或博士学位）。
```
     上述示例中：第一行以 `学历要求：(a)…` 开头；`(b)`、`(c)` 各占新行；前两行以 `；或` 结尾，末行以 `。` 结尾。
   - 若无明确学历条款，填 JSON `null` 或简短说明「原文未单独列出」。

4. **语言要求**（键名：`语言要求`）
   - **仅从** `full_description_block` 中提取**强制性**语言要求（Required / must / fluency 等）。
   - 若**仅**要求英语：统一写 **「英语口语和书面表达流利」**。
   - 若多种语言均为**必要**：写 **「英语及[某语言]口语和书面表达流利」**（按原文语种替换）。
   - **忽略** Desirable / Asset / is desirable / would be an asset 等非强制表述。

5. **时长**（键名：`时长`）
   - **严格依据** `expected_duration` 的原文含义翻译为中文（如「6个月」「2–6个月」「三个月（可延长至六个月）」等），不要臆测。

6. **申请截止时间**（键名：`申请截止时间`）
   - 将 `gray_box_details.deadline` 的英文日期转为 **YYYY年M月D日** 形式（月份与日期可与示例一致不必补零），例如 `Mar 23, 2027` → `2027年3月23日`。

7. **工作地点**（键名：`工作地点`）
   - 以 `work_location` 为主，可结合 `gray_box_details.duty_station`。
   - 城市译为中文惯用地名；若为 Remote / Home-based / Telecommuting 等，译为 **「远程」** 或 **「居家办公」** 等贴切说法；Hybrid 译为 **「混合式」**；若强调灵活办公可用 **「灵活」**。
   - 简短、可读。

8. **详情请见**（键名：`详情请见`）
   - 原样填入原始数据中的 `original_url` 字符串。

## 输出 JSON 的键名（必须完全一致，共 8 个）
`职位`, `组织`, `学历要求`, `语言要求`, `时长`, `申请截止时间`, `工作地点`, `详情请见`

## 重要
- 只输出 JSON 字符串本身，不要使用 markdown 代码围栏。
- 字符串内的引号必须正确转义（JSON 合法）。
- 「学历要求」为多行时，必须在 JSON 字符串中使用 `\\n` 写入换行，以便下游写入 Word 时按行分段显示。
"""


def _require_api_key() -> str:
    key = (GEMINI_API_KEY or "").strip()
    if not key:
        raise RuntimeError(
            "未检测到有效的 GEMINI_API_KEY。\n"
            "请在项目目录创建 .env 并写入一行：GEMINI_API_KEY=你的密钥\n"
            "或在当前会话设置环境变量，例如：\n"
            "  Windows PowerShell: $env:GEMINI_API_KEY=\"你的密钥\""
        )
    return key


def _build_http_options() -> types.HttpOptions | None:
    """若已配置 ``ACTIVE_HTTP_PROXY_URL``，则为 httpx 绑定代理；否则返回 None。"""
    if not ACTIVE_HTTP_PROXY_URL:
        return None
    return types.HttpOptions(
        client_args={"proxy": ACTIVE_HTTP_PROXY_URL},
        async_client_args={"proxy": ACTIVE_HTTP_PROXY_URL},
    )


def _build_genai_client(api_key: str) -> genai.Client:
    opts = _build_http_options()
    if opts is not None:
        print(
            f"  [配置] google-genai Client，httpx proxy={ACTIVE_HTTP_PROXY_URL}",
            flush=True,
        )
        return genai.Client(api_key=api_key, http_options=opts)
    print("  [配置] google-genai Client（未配置 HTTP 代理）", flush=True)
    return genai.Client(api_key=api_key)


def _minimal_safety_settings() -> list[types.SafetySetting]:
    """仅针对文本模型支持的核心危害类别将阈值调至最低（避免图像类枚举触发 INVALID_ARGUMENT）。"""
    categories: list[types.HarmCategory] = [
        types.HarmCategory.HARM_CATEGORY_HATE_SPEECH,
        types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT,
        types.HarmCategory.HARM_CATEGORY_HARASSMENT,
        types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT,
    ]
    return [
        types.SafetySetting(category=c, threshold=types.HarmBlockThreshold.OFF)
        for c in categories
    ]


def _strip_json_fence(text: str) -> str:
    t = text.strip()
    m = re.match(r"^```(?:json)?\s*\n?(.*?)\n?```\s*$", t, re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(1).strip()
    return t


def _parse_model_payload(response: types.GenerateContentResponse) -> dict[str, Any]:
    parsed = getattr(response, "parsed", None)
    if isinstance(parsed, dict):
        return parsed
    text = (response.text or "").strip()
    if not text:
        raise ValueError("模型返回空内容（无 parsed 且无文本）")
    cleaned = _strip_json_fence(text)
    return json.loads(cleaned)


def _build_user_prompt(raw_job: dict[str, Any]) -> str:
    raw_str = json.dumps(raw_job, ensure_ascii=False, indent=2)
    return USER_PROMPT_TEMPLATE.format(raw_job_json=raw_str)


def _generate_content_config() -> types.GenerateContentConfig:
    return types.GenerateContentConfig(
        system_instruction=SYSTEM_INSTRUCTION,
        temperature=0.1,
        response_mime_type="application/json",
        response_schema=JOB_OUTPUT_JSON_SCHEMA,
        safety_settings=_minimal_safety_settings(),
    )


def _format_api_error(exc: BaseException) -> str:
    if isinstance(exc, genai_errors.ClientError):
        return f"{type(exc).__name__}({getattr(exc, 'status', '')}): {exc}"
    return f"{type(exc).__name__}: {exc}"


def _call_model_once(
    client: genai.Client,
    *,
    model: str,
    raw_job: dict[str, Any],
) -> dict[str, Any]:
    """单次调用（不做任何重试）；失败由调用方记录。"""
    user_prompt = _build_user_prompt(raw_job)
    print(f"  [调试] model={model}", flush=True)
    response = client.models.generate_content(
        model=model,
        contents=user_prompt,
        config=_generate_content_config(),
    )
    return _parse_model_payload(response)


def _annotate_result(payload: dict[str, Any], processed_by: str) -> dict[str, Any]:
    out = dict(payload)
    out["processed_by"] = processed_by
    return out


def load_input_rows(path: str) -> list[dict[str, Any]]:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, list):
        raise ValueError(f"{path} 根节点应为数组")
    return data


def _parse_cn_deadline_to_datetime(value: Any) -> datetime | None:
    """
    将「YYYY年M月D日」类中文日期解析为 datetime（当天 00:00:00）。
    解析失败返回 None，排序时置于末尾。
    """
    if value is None:
        return None
    s = str(value).strip().translate(_FULLWIDTH_DIGITS)
    if not s or s.lower() in ("null", "none", "-", "—", ""):
        return None
    m = re.match(r"^(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日\s*$", s)
    if m:
        try:
            y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
            return datetime(y, mo, d)
        except ValueError:
            return None
    m2 = re.match(r"^(\d{4})-(\d{1,2})-(\d{1,2})\s*$", s)
    if m2:
        try:
            y, mo, d = int(m2.group(1)), int(m2.group(2)), int(m2.group(3))
            return datetime(y, mo, d)
        except ValueError:
            return None
    return None


def _sort_results_by_deadline_asc(results: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """按「申请截止时间」升序（越早的截止日越靠前）；无法解析的排在最后，相对顺序稳定。"""

    def key_at(i: int) -> tuple[int, datetime, int]:
        dt = _parse_cn_deadline_to_datetime(results[i].get("申请截止时间"))
        if dt is None:
            return (1, datetime.max, i)
        return (0, dt, i)

    order = sorted(range(len(results)), key=key_at)
    return [results[i] for i in order]


def _strip_json_artifacts_text(s: str) -> str:
    """去除正文不应出现的 JSON 痕迹：ASCII 双引号、行尾逗号。"""
    t = s.replace('"', "'")
    lines = []
    for line in t.split("\n"):
        lines.append(line.rstrip().rstrip(",").rstrip("，"))
    return "\n".join(lines)


def _field_value_to_plain(value: Any) -> str:
    if value is None:
        return "—"
    if isinstance(value, (dict, list)):
        return _strip_json_artifacts_text(
            json.dumps(value, ensure_ascii=False, separators=(",", ":"))
        )
    return _strip_json_artifacts_text(str(value))


def _format_education_with_clause_breaks(text: str) -> str:
    """在 (b) / （b）前插入换行，便于 (a)/(b) 分行显示。"""
    if not text:
        return text
    t = text.replace("\r\n", "\n")
    t = re.sub(r"(?<!\n)\s*((?:\(|（)b(?:\)|）))", r"\n\1", t, flags=re.IGNORECASE)
    return t.strip()


def _strip_education_value_label_prefix_for_docx(text: str) -> str:
    """
    AI 输出的「学历要求」值按规定以「学历要求：」开头；Word 段落已单独渲染键名标签，
    此处去掉正文内重复前缀，保留后续换行与 (a)(b)(c) 结构，避免版式重复。
    """
    t = text.replace("\r\n", "\n")
    if t.startswith("学历要求："):
        return t[len("学历要求：") :].lstrip("\n")
    return t


def _export_weekly_docx(results: list[dict[str, Any]], path: str) -> bool:
    """
    将成功条目导出为 Word；若未安装 python-docx 则打印提示并返回 False。
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt
        from docx.text.run import Run
    except ImportError:
        print(
            "未检测到 python-docx，无法生成 Word。\n"
            "请执行：pip install python-docx\n",
            file=sys.stderr,
            flush=True,
        )
        return False

    def _set_run_font_simsun(run: Run) -> None:
        run.font.name = "SimSun"
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), "SimSun")

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "SimSun"
    normal_r_pr = normal_style._element.get_or_add_rPr()
    normal_r_fonts = normal_r_pr.get_or_add_rFonts()
    normal_r_fonts.set(qn("w:eastAsia"), "SimSun")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("联合国实习岗位半月汇总")
    tr.bold = True
    tr.font.size = Pt(18)
    _set_run_font_simsun(tr)
    doc.add_paragraph()

    for idx, row in enumerate(results):
        if idx > 0:
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(10)
            sep.paragraph_format.space_after = Pt(10)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sep.add_run("──────────────")
            sr.font.size = Pt(9)
            _set_run_font_simsun(sr)
            doc.add_paragraph()

        for key in DOCX_FIELD_KEYS:
            raw_val = row.get(key)
            body = _field_value_to_plain(raw_val)
            if key == "学历要求":
                body = _strip_education_value_label_prefix_for_docx(body)
                body = _format_education_with_clause_breaks(body)

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            label = p.add_run(f"{key}：")
            label.bold = True
            _set_run_font_simsun(label)
            for li, line in enumerate(body.split("\n")):
                if li:
                    brk_run = p.add_run()
                    _set_run_font_simsun(brk_run)
                    brk_run.add_break()
                body_run = p.add_run(line if line else " ")
                _set_run_font_simsun(body_run)

    doc.save(path)
    return True


def _configure_stdio_utf8() -> None:
    try:
        if hasattr(sys.stdout, "reconfigure"):
            sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        if hasattr(sys.stderr, "reconfigure"):
            sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass


def _run_ai_pipeline_blocking(
    rows: list[dict[str, Any]],
    out_json_path: str,
    out_docx_path: str,
    *,
    meta_input_label: str,
) -> dict[str, Any]:
    """同步执行 Gemini 阶梯处理、排序与导出（供 ``asyncio.to_thread`` 调用）。"""
    configure_runtime_networking()
    api_key = _require_api_key()

    out_by_index: dict[int, dict[str, Any]] = {}
    failures: list[dict[str, Any]] = []

    total = len(rows)
    config_snapshot = _generate_content_config()

    with _build_genai_client(api_key) as client:
        for chunk_start in range(0, total, PHASE1_CHUNK_SIZE):
            chunk_end = min(chunk_start + PHASE1_CHUNK_SIZE, total)
            chunk_indices = range(chunk_start, chunk_end)

            print(
                f"\n======== 正在进行第一阶段：Flash 模型（条目 {chunk_start + 1}–{chunk_end} / {total}）========",
                flush=True,
            )

            phase1_failures: list[tuple[int, dict[str, Any], str]] = []

            for i in chunk_indices:
                raw = rows[i]
                url = raw.get("original_url") if isinstance(raw, dict) else ""
                url = url or ""
                print(f"\n[{i + 1}/{total}] 处理: {str(url)[:80]}…", flush=True)

                if not isinstance(raw, dict):
                    msg = "条目不是对象"
                    print(f"  跳过 — {msg}", flush=True)
                    failures.append({"index": i, "phase": 1, "error": msg, "raw": raw})
                    if i < total - 1:
                        time.sleep(THROTTLE_SECONDS)
                    continue

                if raw.get("error"):
                    msg = f"源数据含 error 字段，跳过 API：{raw.get('error')}"
                    print(f"  跳过 — {msg}", flush=True)
                    failures.append(
                        {
                            "index": i,
                            "phase": 1,
                            "original_url": url,
                            "error": msg,
                            "skipped_api": True,
                        }
                    )
                    if i < total - 1:
                        time.sleep(THROTTLE_SECONDS)
                    continue

                try:
                    parsed = _call_model_once(client, model=MODEL_FLASH, raw_job=raw)
                    out_by_index[i] = _annotate_result(parsed, MODEL_FLASH)
                    print("  成功：已解析 JSON（Flash）。", flush=True)
                except Exception as e:
                    err = _format_api_error(e)
                    print(f"  失败（不重试）：{err}", flush=True)
                    traceback.print_exc()
                    phase1_failures.append((i, raw, err))

                if i < total - 1:
                    time.sleep(THROTTLE_SECONDS)

            if phase1_failures:
                print(
                    f"\n======== 正在进行第二阶段：Flash-Lite 模型（补救本批失败 {len(phase1_failures)} 条）========",
                    flush=True,
                )
                # 若本批最后一条已是文件末尾，第一阶段循环不会在末尾 sleep，这里补一次节流
                if chunk_end == total:
                    time.sleep(THROTTLE_SECONDS)
            else:
                print("\n本批第一阶段无失败项，跳过第二阶段。", flush=True)

            for j, (i, raw, phase1_err) in enumerate(phase1_failures):
                url = raw.get("original_url") or ""
                print(
                    f"\n[补救 {j + 1}/{len(phase1_failures)}] 索引 {i + 1}: {str(url)[:80]}…",
                    flush=True,
                )
                print(f"  第一阶段错误：{phase1_err}", flush=True)
                try:
                    parsed = _call_model_once(
                        client, model=MODEL_FLASH_LITE, raw_job=raw
                    )
                    out_by_index[i] = _annotate_result(parsed, MODEL_FLASH_LITE)
                    print("  成功：已解析 JSON（Flash-Lite）。", flush=True)
                except Exception as e:
                    err = _format_api_error(e)
                    print(f"  补救失败：{err}", flush=True)
                    traceback.print_exc()
                    failures.append(
                        {
                            "index": i,
                            "phase": 2,
                            "original_url": url,
                            "phase1_error": phase1_err,
                            "error": err,
                        }
                    )

                if j < len(phase1_failures) - 1 or chunk_end < total:
                    # 节流：补救调用之间、以及批次末尾与下一批之间均等待
                    time.sleep(THROTTLE_SECONDS)

    ordered_results = [out_by_index[k] for k in sorted(out_by_index)]

    print("\n正在排序数据并生成 Word 文档…", flush=True)
    sorted_results = _sort_results_by_deadline_asc(ordered_results)
    print(
        f"  已按「申请截止时间」升序排序（无法解析的日期已置于末尾），"
        f"共 {len(sorted_results)} 条。",
        flush=True,
    )

    payload = {
        "results": sorted_results,
        "failures": failures,
        "meta": {
            "sdk": "google-genai",
            "phase1_model": MODEL_FLASH,
            "phase2_model": MODEL_FLASH_LITE,
            "phase1_chunk_size": PHASE1_CHUNK_SIZE,
            "proxy": ACTIVE_HTTP_PROXY_URL or "none",
            "throttle_seconds": THROTTLE_SECONDS,
            "temperature": config_snapshot.temperature,
            "input": meta_input_label,
            "processed_count": len(sorted_results),
            "failed_count": len(failures),
            "sorted_by": "申请截止时间",
            "sort_order": "asc",
            "docx_output": out_docx_path,
        },
    }

    with open(out_json_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=4)
    print(
        f"  后台备份 JSON 已写入：{out_json_path}（UTF-8，indent=4）。",
        flush=True,
    )

    print(f"  正在写入 Word：{out_docx_path} …", flush=True)
    docx_ok = _export_weekly_docx(sorted_results, out_docx_path)
    if docx_ok:
        print(f"  Word 文档已保存：{out_docx_path}", flush=True)

    print(
        f"\n完成：成功 {len(sorted_results)} 条，失败 {len(failures)} 条。",
        flush=True,
    )
    return {
        "results": sorted_results,
        "failures": failures,
        "payload": payload,
        "output_json_path": out_json_path,
        "output_docx_path": out_docx_path,
        "docx_written": docx_ok,
    }


async def run_ai_processing(
    data: list[dict[str, Any]],
    *,
    output_json_path: str | None = None,
    output_docx_path: str | None = None,
    meta_input_label: str = "pipeline",
) -> dict[str, Any]:
    """
    异步入口：对详情字典列表执行 AI 清洗与 Word 导出（阻塞 SDK 调用置于线程池）。
    """
    oj = output_json_path or OUTPUT_JSON_PATH
    od = output_docx_path or OUTPUT_DOCX_PATH
    return await asyncio.to_thread(
        _run_ai_pipeline_blocking,
        data,
        oj,
        od,
        meta_input_label=meta_input_label,
    )


async def _cli_entry_async() -> None:
    """单独运行 ``ai_processor.py`` 时的 CLI（从 ``INPUT_JSON_PATH`` 读入）。"""
    _configure_stdio_utf8()
    configure_runtime_networking()
    _warn_if_missing_dotenv_file()
    rows = load_input_rows(INPUT_JSON_PATH)
    await run_ai_processing(
        rows,
        output_json_path=OUTPUT_JSON_PATH,
        output_docx_path=OUTPUT_DOCX_PATH,
        meta_input_label=INPUT_JSON_PATH,
    )


def _warn_if_missing_dotenv_file() -> None:
    env_file = os.path.join(os.getcwd(), ".env")
    if os.path.isfile(env_file):
        return
    print(
        "提示：当前工作目录下未找到 .env 文件。\n"
        "建议在项目根目录创建 .env，并添加一行（不要加引号）：\n"
        "  GEMINI_API_KEY=你的密钥\n"
        "若已通过系统环境变量等方式提供 GEMINI_API_KEY，可忽略本提示。\n",
        flush=True,
    )


if __name__ == "__main__":
    try:
        asyncio.run(_cli_entry_async())
    except RuntimeError as e:
        print(str(e), file=sys.stderr)
        sys.exit(1)
    except FileNotFoundError as e:
        print(str(e), file=sys.stderr)
        sys.exit(1)
